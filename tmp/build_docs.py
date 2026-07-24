import os
import re
import subprocess
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

base_dir = r"c:\Users\galih\Documents\Projects\Game\My project"
sub_dir = os.path.join(base_dir, "docs", "submissions")

tugas10_md = """# LAPORAN PROGRESS PROJECT BASED 2 (TUGAS 10)

**Mata Kuliah:** CIE 725 - Game Development  
**Sesi / Pertemuan:** 10  
**Judul Project Game:** Life on Land — Top-Down Tactical Eco-Restoration Simulator (Build Demo)  
**Dosen Pengampu:** 7174 - Ir. Sawali Wahyu, S.Kom., M.Kom  
**Tim Pengembang (Kelompok 9):**  
- 20230801245 — Galih Adhi Kusuma (Lead Programmer)  
- 20230801438 — Firschanya Alula R. (Art Director & Narrative)  
- 20230801205 — Defanda Yeremia C. R. (System Analyst & QA)  
**Progress Project:** 75% (Tahap Prototype Lanjutan & Integrasi PlayFab)

---

## 1. IDENTITAS PROJECT DAN ABSTRAK

- **Nama Game:** Life on Land
- **Genre:** Top-Down Cozy Stage-Based Forest Ecosystem Simulator
- **Platform:** PC / Windows / macOS / Web (Unity Engine 6 - C#)
- **Perspektif:** Fixed 2D Orthographic Top-Down (32 PPU, Pixel Art)
- **Ringkasan Progress (75%):** Pada tahap Project Based 2 ini, game *Life on Land* telah berhasil dikembangkan hingga tahap prototype lanjutan (75% selesai) dengan fokus utama pada **Stage 1: Red Region (The Arid Oasis)**. Sistem mekanik utama seperti pergerakan karakter top-down WASD/Panah, sistem 6 slot hotbar, purifikasi tanah 2 tahap (sekop lalu siram), matriks kelembapan tanah (*soil moisture decay*), serta siklus pertumbuhan tanaman (*Seed -> Sprout -> Mature -> Withered -> Revive*) telah 100% berfungsi di Unity Engine. Integrasi PlayFab Game Manager dan serialisasi data lokal (`SaveData.json`) telah mencapai 50% untuk autentikasi user (Sign-In/Register) dan sinkronisasi progress stage.

---

## 2. KONSEP GAME DAN METODE GAMIFIKASI (STAGE 1 DEMO)

### 2.1 Model Pembelajaran dan Gamifikasi
Game *Life on Land* mengusung model pembelajaran ekologi dan restorasi lingkungan (*Eco-Restoration*) berbasis gamifikasi *Quest-Driven Ecosystem Loop* (`Challenge -> Action -> Reward -> Environmental Shift`). Pemain diajarkan dampak degradasi tanah, pentingnya retensi air, serta keseimbangan kadar oksigen (O2) atmosfer melalui mekanisme permainan yang interaktif.

Tingkatan Gamifikasi (Stage 1 Red Oasis Demo):

1. **Stage 1 — Red Region (The Arid Oasis):**
   - **Aesthetic:** Tanah liat merah, pasir kering, tunggul merah mati.
   - **Tujuan Ekologi:** Mengembalikan tingkat Oksigen lokal dari 15.0% menjadi 50.0%.
   - **NPC:** Maliz (Bear Wrath Barbarian) — Karakter berotot tangguh yang sedih karena oasis terbakarnya.
   - **NPC Mini-Quest:** Mengambil 10 unit air dari kolam dalam (*Deep Pond*) menggunakan ember/gembor air.
   - **Reward:** Desert Shrub Seeds (Tanaman Tipe B — konsumsi air rendah, meningkatkan retensi kelembapan tanah di sekitarnya).
   - **Main Quest:** Membersihkan 5 ubin terbakar (*corrupted tiles*), menanam dan menumbuhkan 5 Semak Gurun (*Desert Shrubs*) hingga mature state, mencapai O2 50.0%.

2. **Rencana Ekspansi Stage (Peta Jalan Masa Depan):**
   - **Stage 2 (Orange Region - Scorched Grove):** Implementasi Pohon Pinus (Tanaman Tipe A — O2 tinggi, butuh air tinggi) dan infrastruktur *Soil Purifier*.
   - **Stage 3 (Pink Bloom - Boss Stage):** Implementasi *Silkmoth Fern* (Tanaman Tipe C — tahan panas), infrastruktur *Irrigation Pipes*, bencana *Heatwave*, dan penyelesaian *Biosphere Dome* untuk menangkap Villain.

---

## 3. ASET GAME DAN SPESIFIKASI MULTIMEDIA (PROGRESS 70-80%)

### 3.1 Karakter dan Aset Visual
- **Protagonis (Umbra):** Karakter utama (*Sloth Monk / Restorer*) dibuat sendiri 100% menggunakan sprite sheet 2D pixel art dengan animasi 4 arah (*Idle, Walk, Dig, Water, Plant*).
- **NPC dan Antagonis:** Sprite Maliz (Bear) dan Villain (Blaze) dirancang khusus dengan gaya pixel art 32 PPU.
- **Aset Lingkungan (Tileset):** Kombinasi aset buatan sendiri (ubin tanah terbakar, ubin *dug soil*, ubin *normal soil*) dan aset pendukung (<50%) untuk variasi pohon mati dan vegetasi dekoratif.
- **Visual Shader Transition:** Menggunakan custom Unity Sprite Shader untuk mentransisikan warna lingkungan dari nuansa kecokelatan/terbakar menjadi hijau cerah secara dinamis seiring peningkatan O2 dari 15% ke 50%.

### 3.2 Background Musik dan Effect (BGM & SFX)
- **BGM Stage 1:** Musik melankolis padang pasir dengan sentuhan alat musik petik.
- **SFX:** Efek suara langkah kaki pada tanah kering, suara cangkul menyentuh tanah terbakar, cipratan air disiramkan, dan efek suara penanaman benih.

---

## 4. ARSITEKTUR KODE DAN KONEKSI DATABASE / PLAYFAB (PROGRESS 50%)

### 4.1 Teknologi dan Bahasa Pemrograman
- **Engine:** Unity 6 (2D Render Pipeline)
- **Bahasa Pemrograman:** C# (.NET Standard 2.1)
- **Format Persistence:** Local JSON Serialization (`SaveData.json`) dan SQLite Database backend.

### 4.2 Integrasi PlayFab Game Manager (Progress 50%)
Untuk memenuhi ketentuan pengkoneksian database/PlayFab:
1. **PlayFab User Authentication:** Fitur Sign-In / Register menggunakan PlayFab Auth (Custom ID & Email Login) telah aktif.
2. **Cloud Save dan Player Title Data:** Data status progress pemain (Stage aktif, jumlah O2 terkumpul, status inventaris hotbar) disinkronkan dari struct `SaveData` ke PlayFab User Data API.
3. **Leaderboard Structure:** Menyiapkan struktur leaderboard nilai restorasi O2 global.

```csharp
using PlayFab;
using PlayFab.ClientModels;
using UnityEngine;

public class PlayFabAuthManager : MonoBehaviour {
    public void LoginCustomID(string customId) {
        var request = new LoginWithCustomIDRequest {
            CustomId = customId,
            CreateAccount = true
        };
        PlayFabClientAPI.LoginWithCustomID(request, OnSuccess, OnError);
    }
    private void OnSuccess(LoginResult result) {
        Debug.Log("PlayFab Login Berhasil! ID: " + result.PlayFabId);
    }
    private void OnError(PlayFabError error) {
        Debug.LogError("PlayFab Login Gagal: " + error.GenerateErrorReport());
    }
}
```

---

## 5. TAMPILAN INTERFACE DAN SCREENSHOT PROTOTYPE (PROGRESS 75%)

Berikut adalah dokumentasi tampilan antarmuka dan visualisasi fitur game *Life on Land* yang telah diimplementasikan di Unity:

### 1. Tampilan Main Menu dan PlayFab Auth
![Tampilan Main Menu Game Life on Land dengan tombol Start, Options, PlayFab Login Panel, dan Logo Game](Assets/Screenshots/main menu.png)  
*Keterangan:* Tampilan menu utama game dengan latar belakang animasi hutan yang perlahan pulih dan form login PlayFab.

### 2. Gameplay Stage 1 — Hotbar dan Purifikasi Tanah
![Gameplay Stage 1 Red Region memperlihatkan Umbra dengan 6 slot Hotbar, ubin burnt, ubin dug, dan status O2 HUD](Assets/Screenshots/grown_trees.png)  
*Keterangan:* Antarmuka permainan top-down memperlihatkan slot hotbar 1-6, bar stamina, indikator O2 (15.0%), dan proses purifikasi tanah menggunakan sekop dan gembor air.

### 3. Tampilan Dialog NPC (Maliz the Bear)
![UI Dialog Box bawah layar menampilkan portrait Maliz Bear dan percakapan quest air](Assets/Screenshots/maliz_dialogs.png)  
*Keterangan:* Sistem dialog visual novel di bagian bawah layar saat Maliz memberikan quest mengambil 10 unit air dari kolam.

### 4. Tampilan Pertumbuhan Tanaman dan Perubahan Visual Environment
![Vegetasi Desert Shrub yang tumbuh dari Sprout ke Mature, serta perubahan warna tanah di sekitarnya](Assets/Screenshots/trees_growing_1.png)  
*Keterangan:* Perubahan state FSM tanaman Semak Gurun dan efek perbaikan lingkungan visual dari warna tanah liat merah menjadi area yang segar.

---

## 6. RENCANA PENYELESAIAN MENUJU 100% DEMO BUILD (NEXT STEPS)

1. Menyelesaikan polis visual efek perubahan O2 di Stage 1 Red Region.
2. Finalisasi quest pengambilan air Maliz the Bear dan penyerahan benih Desert Shrub.
3. Penyempurnaan sinkronisasi data PlayFab Cloud Save untuk pencapaian Achievements dan Leaderboard Global.
4. Polis grafis visual shader dan efek audio responsif.
"""

tugas11_md = """# LAPORAN FINAL AKHIR APLIKASI GAME PROJECT BASED 3 (TUGAS 11)

**Mata Kuliah:** CIE 725 - Game Development  
**Sesi / Pertemuan:** 11  
**Judul Project Game:** Life on Land — Top-Down Tactical Eco-Restoration Simulator (Versi Demo Stage 1)  
**Dosen Pengampu:** 7174 - Ir. Sawali Wahyu, S.Kom., M.Kom  
**Tim Pengembang (Kelompok 9):**  
- 20230801245 — Galih Adhi Kusuma (Lead Programmer)  
- 20230801438 — Firschanya Alula R. (Art Director & Narrative)  
- 20230801205 — Defanda Yeremia C. R. (System Analyst & QA)  
**Progress Project:** 100% SELESAI (Versi Build Demo Playable - Stage 1 Red Region)

---

## 1. PERNYATAAN KETENTUAN FINALISASI GAME DEMO (PROGRESS 100%)

Aplikasi game *Life on Land* telah **100% selesai dikembangkan untuk versi Demo Build** menggunakan Unity Engine (C#). Pengujian dan fungsionalitas game difokuskan secara penuh pada **Stage 1: Red Region (The Arid Oasis)** sebagai inti permainan simulator restorasi ekosistem. Seluruh fitur utama (pergerakan top-down, hotbar 6 slot, purifikasi 2 tahap, siklus kelembapan tanah, FSM pertumbuhan tanaman, dialog NPC Maliz, sistem O2 & stamina, hingga integrasi database/PlayFab 100%) telah diimplementasikan sepenuhnya dan berfungsi 100% tanpa bug.

---

## 2. DOKUMENTASI FITUR-FITUR UTAMA YANG DIIMPLEMENTASIKAN (FITUR VERSI DEMO)

### 2.1 Mekanik Purifikasi Tanah 2-Tahap (Two-Step Purification)
- **Step 1 (Shovel/Sekop):** Pemain menunjuk ubin terbakar gersang (*corrupted burnt tile*) dan menggunakan sekop (Hotbar 1). Ubin berubah state menjadi *Dug Burnt Soil* dengan mengonsumsi 5 Stamina.
- **Step 2 (Watering Can/Gembor Air):** Pemain menyiramkan air dari stok inventaris ke ubin *Dug Burnt Soil* (Hotbar 2). Ubin terpurifikasi penuh menjadi *Normal Clean Soil* yang siap ditanami benih.

### 2.2 Siklus Pertumbuhan Tanaman dan Matriks Kelembapan Tanah (Soil Moisture & Plant FSM)
- **Matriks Kelembapan Tanah (Soil Moisture Grid):** Setiap ubin memiliki nilai moisture (0.0 hingga 1.0) yang menyusut secara bertahap akibat penguapan alami.
- **FSM Pertumbuhan Tanaman:** Tanaman Semak Gurun (*Desert Shrub* - Tipe B) memiliki 4 state pertumbuhan: `Seed` -> `Sprout` -> `Mature Tree` -> `Withered`.
- **Fitur Revive Tanaman Layu:** Tanaman yang layu (*Withered*) akibat kehabisan air tidak mati permanen. Pemain cukup menyiram air langsung ke tanaman atau ubinnya untuk memulihkan (*Revive*) tanaman kembali ke fase Mature.

### 2.3 Sistem Oksigen Atmosfer dan Manajemen Stamina (O2 Buffer & Stamina Debuff)
- **Kalkulasi O2 Dinamis:** Oksigen diawali dari kadar kritis 15.0%. Setiap Semak Gurun mature yang ditanam memancarkan O2 lokal hingga mencapai target Stage 1 yaitu 50.0%.
- **Area-Grid Speed Debuff:** Saat pemain berada di area dengan O2 rendah (< 18.0%), pergerakan karakter Umbra mengalami efek debuff (kecepatan jalan melambat).
- **Stamina & Buffer Recovery:** Berada di zona O2 tinggi (> 18.0%) atau mengonsumsi item air dari kolam akan memulihkan stamina dan buffer oksigen pemain.

### 2.4 Quest NPC Maliz the Bear dan Sistem Dialog Visual Novel
- **NPC Maliz the Bear:** Beruang besar berpenampilan tangguh yang sedih karena oasisnya terbakar.
- **Quest Pengambilan Air (Mini-Quest):** Maliz meminta pemain mengambil 10 unit air dari kolam air dalam (*Deep Pond*). Setelah diserahkan, pemain mendapatkan reward benih Semak Gurun (*Desert Shrub Seeds*).
- **Dialogue UI System:** Antarmuka dialog gaya visual novel di bagian bawah layar lengkap dengan potret karakter (Umbra, Maliz, Villain).

### 2.5 Aset Visual Pixel Art dan Custom Dynamic Shader
- **Karakter Kustom (Umbra & Maliz):** Spritesheet piksel 2D (32 PPU, Point Filter) buatan sendiri dengan animasi 4 arah.
- **Dynamic Post-Processing Shader:** Shader khusus Unity yang mentransisikan atmosfer lingkungan dari warna merah/abu gersang menjadi hijau segar cerah secara mulus seiring meningkatnya persentase O2.
- **BGM & SFX:** Background music melankolis padang pasir, dipadu SFX sekop tanah, cipratan air, dan efek suara penanaman benih.

---

## 3. INTEGRASI DATABASE DAN PLAYFAB MANAGER (100% TERKONEKSI)

Pengembangan game *Life on Land* (Versi Demo) diintegrasikan 100% dengan **PlayFab Game Manager (Cloud API)** dan **Local SQLite/JSON Persistence**.

### 3.1 Skema Data Persistence (SaveData.json / PlayFab User Data)
Data kemajuan permainan tersimpan secara terstruktur dalam format JSON:

```json
{
  "save_meta": {
    "timestamp": "2026-07-22T00:23:46Z",
    "current_level": 1,
    "demo_version": "1.0.0-release",
    "play_time_seconds": 620
  },
  "environment_state": {
    "global_o2_percentage": 50.0,
    "absolute_active_tree_count": 5,
    "global_soil_quality_mean": 0.92
  },
  "player_state": {
    "player_name": "Umbra",
    "current_stamina": 100.0,
    "local_o2_buffer": 100.0,
    "current_water_inventory": 10,
    "active_hotbar_slot": 0,
    "inventory": [
      { "item_id": "desert_shrub_seed", "quantity": 5 }
    ]
  },
  "instantiated_objects": [
    {
      "object_id": "shrub_001",
      "tree_type_id": "desert_shrub",
      "grid_x": 10,
      "grid_y": 14,
      "growth_state": "Mature",
      "ticks_since_last_watered": 0
    }
  ]
}
```

### 3.2 Fitur PlayFab yang Diimplementasikan (100% Connected)
1. **PlayFab User Authentication:** Fitur Sign-In / Register ID pemain, penyimpanan nickname, dan sesi login otomatis.
2. **PlayFab Cloud Save (Title Data & User Progress):** Menyimpan status ubin, inventaris air, dan progress quest Maliz ke cloud PlayFab secara otomatis.
3. **PlayFab Achievements System:**
   - Pencapaian *First Steps* (Menyiram ubin terbakar pertama kali).
   - Pencapaian *Water Bearer* (Menyerahkan 10 unit air ke Maliz).
   - Pencapaian *Green Oasis* (Mencapai O2 50.0% dan memulihkan Red Region).
4. **PlayFab Global Leaderboard:** Peringkat dunia berdasarkan waktu tercepat menyelesaikan restorasi Demo Stage 1.

---

## 4. FLOW TAMPILAN AKHIR GAME DEMO (100% SELESAI)

Berikut adalah dokumentasi alur penuh dari permainan Demo Stage 1 game *Life on Land*:

### 1. Main Menu dan Login PlayFab
![Halaman Main Menu dengan tombol Start Demo, Options, Leaderboard, Achievements, dan Login PlayFab Status Connected](Assets/Screenshots/main menu.png)  
*Keterangan:* Menu utama game memperlihatkan status koneksi PlayFab aktif dan opsi memulai Demo Stage 1.

### 2. Opening Scene — Pembakaran Oasis oleh Antagonis
![Antagonis Villain membakar pohon oasis dan melarikan diri, meninggalkan Maliz yang bersedih](Assets/Screenshots/viallin_dialogs_2.png)  
*Keterangan:* Cutscene awal pembuka demo yang memperkenalkan ancaman pembakaran ekosistem.

### 3. Dialog Quest Maliz the Bear (Fetch Water)
![UI Dialog Box menampilkan Maliz memberikan quest mengambil 10 unit air dari kolam](Assets/Screenshots/maliz_dialogs.png)  
*Keterangan:* Percakapan interaktif Umbra dan Maliz untuk memulai quest pengumpulan air.

### 4. Eksekusi Purifikasi Tanah dan Penyiraman Air
![Umbra mensekop ubin terbakar menjadi dug soil dan menyiramnya menjadi normal clean soil](Assets/Screenshots/watered_soil_1.png)  
*Keterangan:* Demonstrasi fitur mekanik purifikasi tanah 2 tahap di dekat area oasis.

### 5. Penanaman Desert Shrub dan Pertumbuhan FSM
![Semak Gurun tumbuh dari fase Sprout menjadi Mature, dan indikator O2 pada HUD naik drastis](Assets/Screenshots/grown_trees.png)  
*Keterangan:* Proses penanaman benih dan respon peningkatan kadar O2 atmosfer.

### 6. Stage Completion dan Output PlayFab Leaderboard
![Tampilan Layar STAGE 1 CLEARED, shader lingkungan menjadi hijau segar, dan Popup Leaderboard PlayFab](Assets/Screenshots/stage_complete.png)  
*Keterangan:* Tampilan akhir kemenangan versi demo dan pencatatan skor waktu ke leaderboard cloud PlayFab.

---

## 5. KESIMPULAN

Pengembangan aplikasi game *Life on Land* (Versi Demo Stage 1) telah rampung 100% sesuai dengan target spesifikasi demo. Pengujian Usability Alpha (SUS) menghasilkan skor rata-rata 63.45 (Grade D / OK). Game berhasil menyajikan gameplay simulator restorasi ekosistem yang solid, mengintegrasikan fitur purifikasi tanah 2 tahap, FSM tanaman, dynamic shader, quest Maliz, serta konektivitas backend PlayFab secara sempurna.
"""

# Write markdown files
with open(os.path.join(sub_dir, "TUGAS_10_PROJECT_BASED_2.md"), "w", encoding="utf-8") as f:
    f.write(tugas10_md)

with open(os.path.join(sub_dir, "TUGAS_11_PROJECT_BASED_3.md"), "w", encoding="utf-8") as f:
    f.write(tugas11_md)

print("Updated markdown files.")

# Docx Generator with Images & Formatting
def parse_inline_formatting(paragraph, text):
    text = re.sub(r'\\\((.*?)\\\)', r'\1', text)
    text = text.replace(r'\sum', 'Jumlah ').replace(r'\ge', '>=').replace(r'\le', '<=')
    
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`') and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(160, 40, 40)
        else:
            paragraph.add_run(part)

def md_file_to_docx(md_path, docx_path):
    print(f"Building Word doc: {os.path.basename(docx_path)}...")
    doc = docx.Document()
    
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)
        
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        
        rows_data = []
        for line in table_lines:
            if "|" in line:
                cells = [c.strip() for c in line.split("|")]
                if cells and cells[0] == '':
                    cells = cells[1:]
                if cells and cells[-1] == '':
                    cells = cells[:-1]
                
                if cells and all(re.match(r'^-+$', c.replace(':', '')) for c in cells):
                    continue
                rows_data.append(cells)

        if not rows_data:
            table_lines = []
            return

        num_cols = max(len(r) for r in rows_data)
        table = doc.add_table(rows=len(rows_data), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for r_idx, row_cells in enumerate(rows_data):
            row = table.rows[r_idx]
            for c_idx in range(num_cols):
                cell = row.cells[c_idx]
                val = row_cells[c_idx] if c_idx < len(row_cells) else ""
                
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.05
                
                parse_inline_formatting(p, val)

                if r_idx == 0:
                    shading = parse_xml(r'<w:shd {} w:fill="1B365D"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading)
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        r.font.size = Pt(9.5)
                else:
                    if r_idx % 2 == 1:
                        shading = parse_xml(r'<w:shd {} w:fill="F4F6F9"/>'.format(nsdecls('w')))
                        cell._tc.get_or_add_tcPr().append(shading)
                    for r in p.runs:
                        r.font.size = Pt(9.0)

        table_lines = []

    for line in lines:
        raw_line = line.rstrip('\n')
        stripped = raw_line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                shading = parse_xml(r'<w:shd {} w:fill="F0F4F8"/>'.format(nsdecls('w')))
                p._p.get_or_add_pPr().append(shading)
                run = p.add_run("\n".join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(30, 30, 30)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(raw_line)
            continue

        if "|" in stripped and not stripped.startswith("#"):
            table_lines.append(stripped)
            in_table = True
            continue
        else:
            if in_table:
                flush_table()
                in_table = False

        if stripped in ["---", "***", "___"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            continue

        if not stripped:
            continue

        # Image parsing: ![Caption](path)
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)$', stripped)
        if img_match:
            alt_text = img_match.group(1)
            rel_img_path = img_match.group(2)
            full_img_path = os.path.join(base_dir, rel_img_path.replace("/", "\\"))
            if os.path.exists(full_img_path):
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.paragraph_format.space_before = Pt(8)
                img_p.paragraph_format.space_after = Pt(2)
                run = img_p.add_run()
                run.add_picture(full_img_path, width=Inches(5.8))
            else:
                print(f"Warning: Image not found at {full_img_path}")
            continue

        if stripped.startswith("# "):
            p = doc.add_heading(level=1)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            parse_inline_formatting(p, stripped[2:])
            for r in p.runs:
                r.font.size = Pt(18)
                r.font.bold = True
                r.font.color.rgb = RGBColor(27, 54, 93)
        elif stripped.startswith("## "):
            p = doc.add_heading(level=2)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            parse_inline_formatting(p, stripped[3:])
            for r in p.runs:
                r.font.size = Pt(14)
                r.font.bold = True
                r.font.color.rgb = RGBColor(46, 91, 130)
        elif stripped.startswith("### "):
            p = doc.add_heading(level=3)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            parse_inline_formatting(p, stripped[4:])
            for r in p.runs:
                r.font.size = Pt(12)
                r.font.bold = True
                r.font.color.rgb = RGBColor(60, 60, 60)
        elif stripped.startswith("#### "):
            p = doc.add_heading(level=4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            parse_inline_formatting(p, stripped[5:])
            for r in p.runs:
                r.font.size = Pt(11)
                r.font.bold = True
                r.font.color.rgb = RGBColor(80, 80, 80)

        elif re.match(r'^\s*[\*\-]\s+', stripped):
            content = re.sub(r'^\s*[\*\-]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            parse_inline_formatting(p, content)

        elif re.match(r'^\d+\.\s+', stripped):
            m = re.match(r'^(\d+\.)\s+(.*)$', stripped)
            num_prefix = m.group(1)
            content = m.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            r_num = p.add_run(num_prefix + " ")
            r_num.bold = True
            parse_inline_formatting(p, content)

        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            if stripped.startswith("*Keterangan:*"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(8)
            parse_inline_formatting(p, stripped)

    if in_table:
        flush_table()

    doc.save(docx_path)
    print(f"Generated docx: {docx_path}")

# Build docx files
md_file_to_docx(os.path.join(sub_dir, "TUGAS_10_PROJECT_BASED_2.md"), os.path.join(sub_dir, "TUGAS_10_PROJECT_BASED_2.docx"))
md_file_to_docx(os.path.join(sub_dir, "TUGAS_11_PROJECT_BASED_3.md"), os.path.join(sub_dir, "TUGAS_11_PROJECT_BASED_3.docx"))
md_file_to_docx(os.path.join(sub_dir, "TUGAS_11_PROJECT_BASED_3.md"), os.path.join(sub_dir, "TUGAS_11_PROJECT_BASED_3_REVISED.docx"))

# Convert docx to PDF via PowerShell Word COM
ps_script = f"""
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$files = @(
    "{os.path.join(sub_dir, 'TUGAS_10_PROJECT_BASED_2.docx')}",
    "{os.path.join(sub_dir, 'TUGAS_11_PROJECT_BASED_3.docx')}",
    "{os.path.join(sub_dir, 'TUGAS_11_PROJECT_BASED_3_REVISED.docx')}"
)

foreach ($file in $files) {{
    if (Test-Path $file) {{
        $pdfPath = [System.IO.Path]::ChangeExtension($file, ".pdf")
        Write-Host "Converting $file -> $pdfPath"
        $doc = $word.Documents.Open($file)
        $doc.SaveAs([ref]$pdfPath, [ref]17) # 17 = wdFormatPDF
        $doc.Close()
    }}
}}
$word.Quit()
Write-Host "All conversions completed!"
"""

ps_path = os.path.join(base_dir, "tmp", "convert_pdf.ps1")
with open(ps_path, "w", encoding="utf-8") as f:
    f.write(ps_script)

print("Running PowerShell PDF conversion...")
subprocess.run(["powershell", "-ExecutionPolicy", "Bypass", "-File", ps_path], check=True)
print("All tasks completed successfully!")
