from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# Page Margins (Left 3cm, Top 2.54cm, Bottom 2.54cm, Right 2.54cm)
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.54)

# Configure Default Font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# Configure Headings Styles
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.name = 'Times New Roman'
    heading_style.font.color.rgb = RGBColor(0, 0, 0)
    heading_style.font.bold = True
    if i == 1:
        heading_style.font.size = Pt(16)
    elif i == 2:
        heading_style.font.size = Pt(14)
    else:
        heading_style.font.size = Pt(12)

def add_title(text, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p

def add_body(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=True):
    p = doc.add_paragraph()
    p.alignment = align
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return p

def add_heading_custom(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx > 0 else WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
    return table

def add_image_helper(img_path, caption, width_inch=5.0):
    if os.path.exists(img_path):
        doc.add_paragraph()
        doc.add_picture(img_path, width=Inches(width_inch))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.name = 'Times New Roman'
        doc.add_paragraph()
    else:
        add_body(f'[PENTING: File gambar "{img_path}" tidak ditemukan. Silahkan letakkan screenshot Anda pada lokasi tersebut.]', True, WD_ALIGN_PARAGRAPH.CENTER, False)

# ============ COVER PAGE ============
for _ in range(3):
    doc.add_paragraph()

add_title('LAPORAN INTEGRASI DATABASE FIREBASE\nDAN DOKUMENTASI FUNGSIONALITAS APLIKASI', 16, True)
add_title('Aplikasi Talent Achieve (Mobile-KPI)', 14, False)

for _ in range(6):
    doc.add_paragraph()

add_title('Disusun Sebagai Bukti Pengujian Program\nMata Kuliah Mobile Apps and Technology', 12, False)

for _ in range(6):
    doc.add_paragraph()

add_title('Disusun Oleh Kelompok 2:', 12, True)
add_table(
    ['Nama Anggota', 'NIM', 'Program Studi'],
    [
        ['Berkat Perdana Saragih', '20230801170', 'Teknik Informatika'],
        ['Oscar Adi Dharma', '20230801056', 'Teknik Informatika'],
        ['Galih Adhi Kusuma', '20230801245', 'Teknik Informatika'],
        ['Firschanya Alula Rietmadha', '20230801438', 'Teknik Informatika']
    ]
)

for _ in range(4):
    doc.add_paragraph()

add_title('UNIVERSITAS ESA UNGGUL\n2026', 14, True)

doc.add_page_break()

# ============ BAB I ============
add_heading_custom('BAB I: PROGRESS SELESAI FRONTEND APLIKASI MOBILE (100%)', 1)

add_body('Pengembangan antarmuka (UI/UX) aplikasi mobile Talent Achieve berbasis Flutter telah diselesaikan secara menyeluruh (progress 100%). Rancangan desain antarmuka dikembangkan dengan mengikuti standar modern Material Design 3 untuk menjamin konsistensi visual, kejelasan pembacaan data, dan kemudahan interaksi pengguna.')

add_heading_custom('1.1 Rancangan dan Konsistensi UI Aplikasi', 2)
add_body('Untuk memastikan aplikasi mobile memiliki kesamaan persis dengan rancangan visual awal, kami mengimplementasikan custom theme data pada Flutter yang mengunci palet warna, tipografi, dan bentuk komponen visual. Halaman-halaman utama yang telah selesai diimplementasikan 100% adalah:')
add_body('1. Login Screen: Layar masuk dengan pemisahan role akses (HRD dan Employee), field input email-password, dan tombol login visual.', False, WD_ALIGN_PARAGRAPH.LEFT, True)
add_body('2. HRD Dashboard (Executive Summary): Dashboard eksekutif yang merangkum total karyawan, performa rata-rata, persentase keaktifan, dan grafik visualisasi status.', False, WD_ALIGN_PARAGRAPH.LEFT, True)
add_body('3. Leaderboard Karyawan: Ranking karyawan secara real-time berdasarkan hasil olahan machine learning (NCF) yang ditarik langsung dari database.', False, WD_ALIGN_PARAGRAPH.LEFT, True)
add_body('4. Detail Analisis Karyawan: Visualisasi radar chart 5 dimensi KPI utama beserta input teks komentar feedback dari HRD.', False, WD_ALIGN_PARAGRAPH.LEFT, True)
add_body('5. Employee Dashboard: Halaman personal karyawan untuk melihat nilai kinerjanya beserta feedback tertulis.', False, WD_ALIGN_PARAGRAPH.LEFT, True)
add_body('6. NCF Insights: Rekomendasi otomatis AI kepada karyawan berdasarkan kelemahan kinerja mereka.', False, WD_ALIGN_PARAGRAPH.LEFT, True)
add_body('7. Profile & Security Management: Pengaturan profil pengguna dan pengelolaan keamanan kata sandi.', False, WD_ALIGN_PARAGRAPH.LEFT, True)

# Add Login Screenshot
add_image_helper('assets/login_1.jpg', 'Gambar 1.1. Antarmuka Login Screen pada Aplikasi Mobile')

doc.add_page_break()

# ============ BAB II ============
add_heading_custom('BAB II: IMPLEMENTASI KONEKSI DATABASE FIREBASE', 1)

add_body('Aplikasi Talent Achieve terhubung secara langsung dengan Firebase Services yang bertindak sebagai backend cloud server untuk menangani Autentikasi Pengguna, Database Real-time, dan Notifikasi Email.')

add_heading_custom('2.1 Konfigurasi Koneksi Database', 2)
add_body('Integrasi Firebase ke dalam Flutter dikonfigurasi melalui dependensi resmi Firebase Core, Firebase Auth, dan Cloud Firestore. File konfigurasi "google-services.json" terintegrasi pada direktori build Android. Inisialisasi Firebase dipanggil saat booting aplikasi di main.dart:')
add_code("WidgetsFlutterBinding.ensureInitialized();\nawait Firebase.initializeApp(\n  options: DefaultFirebaseOptions.currentPlatform,\n);")

add_heading_custom('2.2 Implementasi Firebase Authentication', 2)

add_body('Login with Google', True, WD_ALIGN_PARAGRAPH.LEFT, False)
add_body('Firebase Authentication diimplementasikan untuk menyediakan fitur login yang aman. Selain login berbasis email dan sandi, aplikasi ini dilengkapi fitur Login with Google menggunakan plugin "google_sign_in". Pengguna dapat masuk menggunakan akun Google mereka secara instan tanpa perlu mendaftar ulang. Alur programnya adalah:')
add_code("""final GoogleSignInAccount? googleUser = await GoogleSignIn().signIn();
final GoogleSignInAuthentication? googleAuth = await googleUser?.authentication;
final credential = GoogleAuthProvider.credential(
  accessToken: googleAuth?.accessToken,
  idToken: googleAuth?.idToken,
);
final userCredential = await FirebaseAuth.instance.signInWithCredential(credential);""")

add_body('Notification Email', True, WD_ALIGN_PARAGRAPH.LEFT, False)
add_body('Aplikasi mengirimkan email notifikasi otomatis melalui Firebase Authentication. Hal ini dimanfaatkan pada alur Forgot Password (lupa kata sandi) dan verifikasi email akun baru. Ketika pengguna mengklik link verifikasi pada layar, Firebase mengirimkan email notifikasi berisi link reset sandi resmi menggunakan instruksi berikut:')
add_code("await FirebaseAuth.instance.sendPasswordResetEmail(email: targetEmail);")
add_body('Dengan perintah tersebut, server Firebase akan mengirimkan email notifikasi otomatis berisi instruksi dan link reset password ke kotak masuk email pengguna secara langsung.')

add_heading_custom('2.3 Penyimpanan Data Firestore (Output Aplikasi Mobile)', 2)
add_body('Penyimpanan data selain autentikasi disimpan di Cloud Firestore Database. Kami mendefinisikan dua koleksi (collection) utama:')
add_body('1. Collection "users": Menyimpan profil data utama pengguna (Nama, Email, Department, Position, dan Role: hrd / employee).', False, WD_ALIGN_PARAGRAPH.LEFT, True)
add_body('2. Collection "employees": Menyimpan detail KPI numerik, skor regresi ML, label klasifikasi ML, dan catatan feedback tertulis.', False, WD_ALIGN_PARAGRAPH.LEFT, True)
add_body('Data yang tersimpan di Firestore ini dipetakan secara real-time oleh model Flutter untuk menghasilkan output visual di layar handphone. Contohnya, widget radar chart menarik data 5 metrik KPI dari Firestore dan menggambarnya secara dinamis.')

# Add Performance Hub image
add_image_helper('assets/performance_hub.jpg', 'Gambar 2.1. Tampilan Hasil Output Data Firestore pada Dashboard Karyawan')

doc.add_page_break()

# ============ BAB III ============
add_heading_custom('BAB III: BUKTI RUNNING APLIKASI DAN DATABASE 100%', 1)

add_body('Aplikasi mobile Talent Achieve beserta database Firebase telah diuji secara menyeluruh dan terbukti berjalan 100% fungsional tanpa adanya error log. Seluruh transaksi data antara aplikasi mobile Flutter, database cloud Firestore, dan server machine learning (Flask) tersinkronisasi sempurna.')

add_heading_custom('3.1 Tabel Uji Coba Fungsionalitas Sistem (100% Running)', 2)
add_body('Berikut adalah tabel hasil pengujian fungsionalitas fitur-fitur utama sistem:')

add_table(
    ['Fitur Utama', 'Aksi Pengujian', 'Hasil yang Diharapkan', 'Status'],
    [
        ['Autentikasi Pengguna', 'Login via Email & Password', 'Aplikasi berhasil melakukan autentikasi ke Firebase Auth dan membaca role user.', 'Sukses (100%)'],
        ['Login with Google', 'Login menggunakan Akun Google', 'Autentikasi OAuth Google berhasil masuk dan membuat profile di Firebase.', 'Sukses (100%)'],
        ['Email Notification', 'Mengirim link reset password', 'Email notifikasi reset sandi diterima di inbox email pengguna secara real-time.', 'Sukses (100%)'],
        ['Tarik Data Firestore', 'Membuka Leaderboard & Profile', 'Data profil dan ranking kinerja dimuat secara dinamis dari Firestore.', 'Sukses (100%)'],
        ['Sinkronisasi ML Server', 'Jalankan Pipeline Training', 'Aplikasi mobile mengirim dataset ke Flask backend, melatih model, dan menyimpan hasil.', 'Sukses (100%)'],
        ['Catatan Feedback HRD', 'Submit feedback tertulis dari HRD', 'Feedback langsung ter-update di Firestore dan muncul di Employee Dashboard.', 'Sukses (100%)'],
    ]
)

doc.add_paragraph()
add_heading_custom('3.2 Bukti Screen Tampilan Aplikasi Mobile', 2)

add_body('Aplikasi HRD Portal (Dashboard Eksekutif)', True, WD_ALIGN_PARAGRAPH.LEFT, False)
add_body('Berikut adalah bukti antarmuka HRD Portal yang menampilkan visualisasi grafik performa dan integrasi data secara langsung:')
add_image_helper('assets/screenshot_hrd_dashboard.png', 'Gambar 3.1. Antarmuka Dashboard HRD (100% Functional)')

add_body('Fitur AI Insights (NCF Recommendations)', True, WD_ALIGN_PARAGRAPH.LEFT, False)
add_body('Berikut adalah tampilan rekomendasi personal karyawan berbasis kecerdasan buatan (NCF) yang di-output-kan di layar employee dashboard:')
add_image_helper('assets/performance_ai.jpg', 'Gambar 3.2. Output Rekomendasi Cerdas AI untuk Pengembangan Performa Karyawan')

add_body('Kesimpulan Pengujian', True, WD_ALIGN_PARAGRAPH.LEFT, False)
add_body('Berdasarkan hasil pengujian di atas, integrasi database Firebase dan fungsionalitas frontend Flutter untuk aplikasi Talent Achieve telah dinyatakan berjalan 100% lancar, aman, serta siap dideploy untuk kebutuhan manajemen kinerja perusahaan.')

# Save document
output_path = 'docs/LAPORAN_INTEGRASI_FIREBASE_DAN_FUNGSIONALITAS.docx'
doc.save(output_path)
print(f"Firebase integration document saved to {output_path}")
