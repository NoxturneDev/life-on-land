import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_poster():
    sub_dir = r"c:\Users\galih\Documents\Projects\Game\My project\docs\submissions"
    screenshots_dir = r"c:\Users\galih\Documents\Projects\Game\My project\Assets\Screenshots"
    charts_dir = os.path.join(sub_dir, "charts")
    docx_path = os.path.join(sub_dir, "TUGAS_13_POSTER_A4_HORIZONTAL.docx")

    print(f"Generating A4 Horizontal Poster Word document: {docx_path}")
    doc = docx.Document()

    # Section A4 Landscape Setup
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)

    # 1. Header Banner Table
    banner_table = doc.add_table(rows=1, cols=1)
    banner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner_table.autofit = False
    cell = banner_table.rows[0].cells[0]
    cell.width = Inches(10.89)
    set_cell_shading(cell, "1B365D")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)

    p_title = cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_title.add_run("LIFE ON LAND: TOP-DOWN TACTICAL ECO-RESTORATION SIMULATOR\n")
    r_t.font.bold = True
    r_t.font.name = "Arial"
    r_t.font.size = Pt(16)
    r_t.font.color.rgb = RGBColor(255, 255, 255)

    r_s = p_title.add_run("POSTER PROYEK AKHIR GAME DEVELOPMENT — UNIVERSITAS ESA UNGGUL (2026)")
    r_s.font.bold = True
    r_s.font.name = "Arial"
    r_s.font.size = Pt(11)
    r_s.font.color.rgb = RGBColor(74, 185, 120)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 2. Main 3-Column Layout Table
    layout_table = doc.add_table(rows=1, cols=3)
    layout_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    layout_table.autofit = False

    cols = layout_table.rows[0].cells
    widths = [Inches(3.5), Inches(3.8), Inches(3.59)]

    for i, c in enumerate(cols):
        c.width = widths[i]
        set_cell_margins(c, top=80, bottom=80, left=100, right=100)
        set_cell_shading(c, "F8F9FA")

    # --- COLUMN 1: Identitas & Konsep & Metode ---
    c1 = cols[0]
    p = c1.paragraphs[0]
    
    # Block A: Identitas Tim & Dosen
    r = p.add_run("👥 IDENTITAS TIM & DOSEN\n")
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(27, 54, 93)
    
    p_id = c1.add_paragraph()
    p_id.paragraph_format.line_spacing = 1.1
    p_id.paragraph_format.space_after = Pt(6)
    p_id.add_run("Dosen Pengampu:\n").bold = True
    p_id.add_run("Ir. Sawali Wahyu, S.Kom., M.Kom.\n\n")
    p_id.add_run("Anggota Kelompok:\n").bold = True
    p_id.add_run("1. Galih Adhi Kusuma (20230801198)\n   [Lead Programmer & Backend]\n")
    p_id.add_run("2. Firschanya Alula R. (20230801201)\n   [Art Director & Narrative]\n")
    p_id.add_run("3. Defanda Yeremia C. R. (20230801205)\n   [System Analyst & QA Tester]\n")

    # Block B: Abstrak & Konsep
    p_ab = c1.add_paragraph()
    r = p_ab.add_run("🌱 OVERVIEW & KONSEP GAME\n")
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(27, 54, 93)
    
    p_ab_txt = c1.add_paragraph()
    p_ab_txt.paragraph_format.line_spacing = 1.05
    p_ab_txt.paragraph_format.space_after = Pt(6)
    p_ab_txt.add_run(
        "Bumi pasca-apokaliptik mengalami krisis atmosfer dengan kadar oksigen tersisa 15.0%. "
        "Pemain berperan sebagai Restorer terakhir (Umbra) yang memulihkan biosfer tanah demi tanah "
        "sambil mengejar antagonis Blaze melintasi 3 wilayah (Red, Orange, dan Pink Region)."
    )

    # Block C: Tujuan, Fungsi & Manfaat
    p_tf = c1.add_paragraph()
    r = p_tf.add_run("🎯 TUJUAN & MANFAAT\n")
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(27, 54, 93)
    
    p_tf_txt = c1.add_paragraph()
    p_tf_txt.paragraph_format.line_spacing = 1.05
    p_tf_txt.add_run("• Tujuan: ").bold = True
    p_tf_txt.add_run("Menaikkan O2 atmosferik hingga >= 21.0% dan memurnikan tanah terpolusi.\n")
    p_tf_txt.add_run("• Manfaat User: ").bold = True
    p_tf_txt.add_run("Edukasi ekologi & simulasi taktis bercocok tanam.\n")
    p_tf_txt.add_run("• Manfaat Reviewer: ").bold = True
    p_tf_txt.add_run("Verifikasi arsitektur FSM, Grid Matrix & Cloud Save.")

    # --- COLUMN 2: Metode, Rancangan & Screenshots ---
    c2 = cols[1]
    p2 = c2.paragraphs[0]

    # Block D: Metode & Rancangan Engine
    r = p2.add_run("🛠️ METODE & RANCANGAN GAME\n")
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(27, 54, 93)

    p_rd = c2.add_paragraph()
    p_rd.paragraph_format.line_spacing = 1.05
    p_rd.paragraph_format.space_after = Pt(6)
    p_rd.add_run("• Genre & Engine: ").bold = True
    p_rd.add_run("Top-Down Tactical Eco-Restoration Simulator (Unity 6 / C#).\n")
    p_rd.add_run("• Metode Dev: ").bold = True
    p_rd.add_run("Game Development Life Cycle (GDLC - 6 Tahapan Iteratif).\n")
    p_rd.add_run("• Rancangan Sistem: ").bold = True
    p_rd.add_run("FSM daur hidup vegetasi, Grid World Matrix kelembapan tanah, dan Gamification Challenge-Action-Reward-Environmental Shift.\n")

    # Block E: Output Visual Screenshot Game
    r_img = c2.add_paragraph().add_run("🖼️ TAMPILAN OUTPUT GAME\n")
    r_img.font.bold = True; r_img.font.size = Pt(11); r_img.font.color.rgb = RGBColor(27, 54, 93)

    # Add game screenshots
    img1 = os.path.join(screenshots_dir, "maliz_dialogs.png")
    img2 = os.path.join(screenshots_dir, "grown_trees.png")

    if os.path.exists(img1):
        p_i1 = c2.add_paragraph()
        p_i1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_i1.paragraph_format.space_after = Pt(4)
        p_i1.add_run().add_picture(img1, width=Inches(3.4))
        r_cap1 = p_i1.add_run("\nTampilan Gameplay: Purifikasi Tanah & Quest Maliz")
        r_cap1.font.size = Pt(8.5); r_cap1.font.italic = True

    if os.path.exists(img2):
        p_i2 = c2.add_paragraph()
        p_i2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_i2.paragraph_format.space_after = Pt(4)
        p_i2.add_run().add_picture(img2, width=Inches(3.4))
        r_cap2 = p_i2.add_run("\nTampilan Vegetasi Pohon Dewasa & Emisi O2")
        r_cap2.font.size = Pt(8.5); r_cap2.font.italic = True

    # --- COLUMN 3: Hasil Testing & Kesimpulan ---
    c3 = cols[2]
    p3 = c3.paragraphs[0]

    # Block F: Hasil Testing Aplikasi
    r = p3.add_run("📊 HASIL TESTING APLIKASI\n")
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(27, 54, 93)

    p_ts = c3.add_paragraph()
    p_ts.paragraph_format.line_spacing = 1.05
    p_ts.paragraph_format.space_after = Pt(4)
    p_ts.add_run("• Alpha Testing (SUS): ").bold = True
    p_ts.add_run("Rata-rata skor 63.45 (Grade D / Marginal High - OK) dari 21 responden.\n")
    p_ts.add_run("• Beta Testing (UAT): ").bold = True
    p_ts.add_run("Rata-rata persentase keberhasilan 82.4% (Kategori SANGAT LAYAK).\n")

    # Add Chart Image
    chart_img = os.path.join(charts_dir, "gform_chart_5_sus_scores.png")
    if os.path.exists(chart_img):
        p_ci = c3.add_paragraph()
        p_ci.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ci.paragraph_format.space_after = Pt(6)
        p_ci.add_run().add_picture(chart_img, width=Inches(3.2))
        r_ccap = p_ci.add_run("\nGrafik Distribusi Skor Testing Usability SUS")
        r_ccap.font.size = Pt(8.5); r_ccap.font.italic = True

    # Block G: Kesimpulan & Saran
    r_ks = c3.add_paragraph().add_run("📌 KESIMPULAN & SARAN\n")
    r_ks.font.bold = True; r_ks.font.size = Pt(11); r_ks.font.color.rgb = RGBColor(27, 54, 93)

    p_ks_txt = c3.add_paragraph()
    p_ks_txt.paragraph_format.line_spacing = 1.05
    p_ks_txt.add_run("• Kesimpulan: ").bold = True
    p_ks_txt.add_run(
        "Game Life on Land Demo Stage 1 berhasil dibangun 100% menggunakan Unity C# dengan mekanik purifikasi 2-tahap, "
        "FSM daur hidup tanaman, quest air, dan recovery O2 50.0%.\n"
    )
    p_ks_txt.add_run("• Saran: ").bold = True
    p_ks_txt.add_run(
        "Pengembangan Stage 2 (Soil Purifier) & Stage 3 (Pipa Irigasi, Heatwave & penangkapan Blaze), "
        "serta pengimbangan kontrol touch-screen mobile."
    )

    doc.save(docx_path)
    print(f"Successfully created clean A4 Horizontal Poster Word doc: {docx_path}")

if __name__ == "__main__":
    create_poster()
