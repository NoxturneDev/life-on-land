import re
import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def parse_inline_formatting(paragraph, text):
    text = re.sub(r'\\\((.*?)\\\)', r'\1', text)
    text = text.replace(r'\sum', 'Jumlah ').replace(r'\ge', '>=').replace(r'\le', '<=')
    
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`') and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(160, 40, 40)
        else:
            paragraph.add_run(part)

def md_file_to_docx(md_path, docx_path):
    print(f"Building clean Word document: {os.path.basename(docx_path)}...")
    doc = docx.Document()
    
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)
        
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        
        rows_data = []
        for line in table_lines:
            if "|" in line:
                cells = [c.strip() for c in line.split("|")]
                if cells and cells[0] == '':
                    cells = cells[1:]
                if cells and cells[-1] == '':
                    cells = cells[:-1]
                
                if cells and all(re.match(r'^-+$', c.replace(':', '')) for c in cells):
                    continue
                rows_data.append(cells)

        if not rows_data:
            table_lines = []
            return

        num_cols = max(len(r) for r in rows_data)
        table = doc.add_table(rows=len(rows_data), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for r_idx, row_cells in enumerate(rows_data):
            row = table.rows[r_idx]
            for c_idx in range(num_cols):
                cell = row.cells[c_idx]
                val = row_cells[c_idx] if c_idx < len(row_cells) else ""
                
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.05
                
                parse_inline_formatting(p, val)

                if r_idx == 0:
                    shading = parse_xml(r'<w:shd {} w:fill="1B365D"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading)
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        r.font.size = Pt(9.5)
                else:
                    if r_idx % 2 == 1:
                        shading = parse_xml(r'<w:shd {} w:fill="F4F6F9"/>'.format(nsdecls('w')))
                        cell._tc.get_or_add_tcPr().append(shading)
                    for r in p.runs:
                        r.font.size = Pt(9.0)

        table_lines = []

    for line in lines:
        raw_line = line.rstrip('\n')
        stripped = raw_line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                shading = parse_xml(r'<w:shd {} w:fill="F0F4F8"/>'.format(nsdecls('w')))
                p._p.get_or_add_pPr().append(shading)
                run = p.add_run("\n".join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(30, 30, 30)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(raw_line)
            continue

        if "|" in stripped and not stripped.startswith("#"):
            table_lines.append(stripped)
            in_table = True
            continue
        else:
            if in_table:
                flush_table()
                in_table = False

        if stripped in ["---", "***", "___"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            continue

        if not stripped:
            continue

        if stripped.startswith("# "):
            p = doc.add_heading(level=1)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            parse_inline_formatting(p, stripped[2:])
            for r in p.runs:
                r.font.size = Pt(18)
                r.font.bold = True
                r.font.color.rgb = RGBColor(27, 54, 93)
        elif stripped.startswith("## "):
            p = doc.add_heading(level=2)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            parse_inline_formatting(p, stripped[3:])
            for r in p.runs:
                r.font.size = Pt(14)
                r.font.bold = True
                r.font.color.rgb = RGBColor(46, 91, 130)
        elif stripped.startswith("### "):
            p = doc.add_heading(level=3)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            parse_inline_formatting(p, stripped[4:])
            for r in p.runs:
                r.font.size = Pt(12)
                r.font.bold = True
                r.font.color.rgb = RGBColor(60, 60, 60)
        elif stripped.startswith("#### "):
            p = doc.add_heading(level=4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            parse_inline_formatting(p, stripped[5:])
            for r in p.runs:
                r.font.size = Pt(11)
                r.font.bold = True
                r.font.color.rgb = RGBColor(80, 80, 80)

        elif re.match(r'^\s*[\*\-]\s+', stripped):
            content = re.sub(r'^\s*[\*\-]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            parse_inline_formatting(p, content)

        elif re.match(r'^\d+\.\s+', stripped):
            m = re.match(r'^(\d+\.)\s+(.*)$', stripped)
            num_prefix = m.group(1)
            content = m.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            r_num = p.add_run(num_prefix + " ")
            r_num.bold = True
            parse_inline_formatting(p, content)

        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            parse_inline_formatting(p, stripped)

    if in_table:
        flush_table()

    doc.save(docx_path)
    print(f"Successfully generated clean Word doc: {docx_path}")

sub_dir = r"c:\Users\galih\Documents\Projects\Game\My project\docs\submissions"

files_to_convert = [
    ("TUGAS_10_PROJECT_BASED_2.md", "TUGAS_10_PROJECT_BASED_2.docx"),
    ("TUGAS_11_PROJECT_BASED_3.md", "TUGAS_11_PROJECT_BASED_3.docx"),
    ("TUGAS_11_PROJECT_BASED_3.md", "TUGAS_11_PROJECT_BASED_3_REVISED.docx"),
    ("TUGAS_12_PROJECT_BASED_4_TESTING.md", "TUGAS_12_PROJECT_BASED_4_TESTING.docx"),
    ("TUGAS_13_LAPORAN_AKHIR_5_BAB.md", "TUGAS_13_LAPORAN_AKHIR_5_BAB.docx"),
    ("TUGAS_13_PPT_PRESENTASI.md", "TUGAS_13_PPT_PRESENTASI.docx"),
    ("TUGAS_13_POSTER_CONTENT.md", "TUGAS_13_POSTER_CONTENT.docx"),
    ("TUGAS_13_POSTER_A4_HORIZONTAL.md", "TUGAS_13_POSTER_A4_HORIZONTAL.docx"),
    ("TUGAS_13_DOKUMENTASI_TUTORIAL_DAN_TEKNIKAL_GUIDE.md", "TUGAS_13_DOKUMENTASI_TUTORIAL_DAN_TEKNIKAL_GUIDE.docx"),
    ("DOKUMENTASI_TEKNIS_LIFE_ON_LAND.md", "DOKUMENTASI_TEKNIS_LIFE_ON_LAND.docx"),
    ("GUIDELINE_PENGGUNA_LIFE_ON_LAND.md", "GUIDELINE_PENGGUNA_LIFE_ON_LAND.docx"),
]

for md, docx_name in files_to_convert:
    md_p = os.path.join(sub_dir, md)
    docx_p = os.path.join(sub_dir, docx_name)
    if os.path.exists(md_p):
        try:
            md_file_to_docx(md_p, docx_p)
        except Exception as e:
            print(f"Error converting {md}: {e}")

