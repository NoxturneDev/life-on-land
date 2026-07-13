import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

doc = docx.Document()

# Base Setup
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Cover Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = title_p.add_run("LAPORAN DOKUMEN KONSEP, SKENARIO, DAN ARSITEKTUR GAME\n")
run_title.bold = True
run_title.font.size = Pt(16)
run_title.font.color.rgb = RGBColor(34, 112, 63) # Cozy green color

run_sub = title_p.add_run("GAME: \"LIFE ON LAND\"\n(Restorasi Ekologi Pasca-Apokaliptik)")
run_sub.bold = True
run_sub.font.size = Pt(13)
run_sub.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph().add_run("\n" + "="*80 + "\n").bold = True

# 1. Gamification Scenario
doc.add_heading("1. Skenario Konsep Gamifikasi Game \"Life on Land\"", level=1)
p1 = doc.add_paragraph()
p1.add_run("Berikut adalah tabel konsep gamifikasi yang ditawarkan oleh game ")
p1.add_run("\"Life on Land\"").bold = True
p1.add_run(" untuk mengedukasi pemain mengenai pentingnya ekologi, fotosintesis, dan manajemen air:")

table = doc.add_table(rows=1, cols=3)
table.alignment = 1
table.autofit = False

hdr_cells = table.rows[0].cells
headers = ["No", "Elemen Gamifikasi", "Deskripsi Skenario Game"]
widths = [Inches(0.5), Inches(2.0), Inches(4.5)]

for i, title in enumerate(headers):
    hdr_cells[i].text = title
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(hdr_cells[i], "22703F")
    set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
    hdr_cells[i].width = widths[i]

scenarios = [
    ("1", "Tujuan Permainan\n(Game Goal)", 
     "Memulihkan atmosfer bumi pasca-apokaliptik dari kadar oksigen awal 15.0% menjadi zona layak huni >= 21.0%.\nPemain harus menanam dan merawat minimal 1000 pohon dewasa, serta mendirikan bangunan ekologi unik sebagai penanda akhir permainan."),
    
    ("2", "Alur Cerita\n(Storyline)", 
     "Pemain adalah satu-satunya 'Restorer' yang selamat di bumi pasca-apokaliptik gersang.\nDipandu oleh robot pemantau ekologi (drone), pemain harus menjelajah daerah kering, mengumpulkan sisa air dari kolam, dan menanam hutan penyedia oksigen sebelum stamina dan cadangan udaranya habis."),
    
    ("3", "Model Pembelajaran\n(Learning Model)", 
     "Experiential Learning & Problem-Based Learning.\nPemain secara praktis belajar tentang:\n- Siklus air tanah (evaporasi ubin & konsumsi tanaman).\n- Tingkat emisi O2 tanaman berdasarkan siklus tumbuh (Seed -> Sprout -> Mature).\n- Dampak iklim (Disaster ticks: gelombang panas & hama)."),
    
    ("4", "Tingkat Kenyamanan\n(Comfort Level)", 
     "Cozy Simulation / High Comfort.\nTidak ada sistem kematian instan. Kontrol pergerakan (berjalan, melompat, meluncur/dash) terasa responsif dan diiringi dengan musik ambient alam yang tenang guna meningkatkan fokus belajar."),
    
    ("5", "Tingkat Kesulitan\n(Difficulty)", 
     "Escalating Difficulty.\nTantangan bertambah seiring bertambahnya level pemain (setiap kelipatan level 5/10), seperti timbulnya bencana kekeringan gelombang panas (mengurangi 50% kelembaban tanah) dan hama tanaman."),
    
    ("6", "Reward & Leaderboard", 
     "Reward:\n- Pengalaman (XP) dan blueprint bangunan baru (Soil Purifier, Irrigation Pipes).\n- Transisi visual dunia dari cokelat sepia gersang menjadi hutan hijau yang hidup.\nLeaderboard:\n- Pencatatan lokal efisiensi air, jumlah pohon hidup, dan durasi restorasi tercepat.")
]

for no, elem, desc in scenarios:
    row_cells = table.add_row().cells
    row_cells[0].text = no
    row_cells[1].text = elem
    row_cells[2].text = desc
    
    # Bold the element column
    row_cells[1].paragraphs[0].runs[0].font.bold = True
    
    for i in range(3):
        set_cell_margins(row_cells[i], top=100, bottom=100, left=120, right=120)
        row_cells[i].width = widths[i]
        # Alternate row coloring
        if int(no) % 2 == 0:
            set_cell_shading(row_cells[i], "F0F5F2")

doc.add_paragraph("\n" + "_"*50 + "\n")

# 2. Information Architecture & Navigation Structure
doc.add_heading("2. Arsitektur Informasi & Struktur Navigasi", level=1)
p2 = doc.add_paragraph()
p2.add_run("Model struktur navigasi game dirancang menggunakan tipe ")
p2.add_run("Hierarki (Tree Structure)").bold = True
p2.add_run(" untuk mempermudah alur interaksi pemain:")

nav_p = doc.add_paragraph()
nav_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
nav_run = nav_p.add_run(
    "               +-----------------------+\n"
    "               |       MAIN MENU       |\n"
    "               +-----------------------+\n"
    "                   /       |       \\\n"
    "        +---------+  +-----+-----+  +---------+\n"
    "        |  MULAI  |  | PENGATURAN |  |  KELUAR |\n"
    "        +----+----+  +-----+-----+  +---------+\n"
    "             |             | \n"
    "    +--------v--------+    ├── Volume Musik & SFX\n"
    "    | GAMEPLAY SCENE  |    └── Tombol Kontrol\n"
    "    +--------+--------+\n"
    "             | \n"
    "             ├── HUD/UI Pemain (Stamina, Air, Benih, Oksigen)\n"
    "             ├── Interaksi Grid (Menanam, Menyiram, Membangun)\n"
    "             └── Menu Jeda (Pause)\n"
    "                     ├── Lanjutkan\n"
    "                     └── Simpan & Keluar\n"
)
nav_run.font.name = 'Courier New'
nav_run.font.size = Pt(10)
nav_run.bold = True

doc.add_paragraph("\n" + "_"*50 + "\n")

# 3. Application Architecture
doc.add_heading("3. Arsitektur Aplikasi Game Menyeluruh", level=1)
p3 = doc.add_paragraph()
p3.add_run("Arsitektur aplikasi game \"Life on Land\" dibangun secara monolitik lokal untuk menjamin performa tinggi tanpa koneksi internet. Diagram arsitektur berikut digambarkan menggunakan format Mermaid JS:")

mermaid_p = doc.add_paragraph()
mermaid_run = mermaid_p.add_run(
    "flowchart TD\n"
    "    subgraph Client_Application [Client Application]\n"
    "        Unity[\"Unity Engine<br/>Rendering & Physics\"]\n"
    "        Gameplay[\"C# Gameplay Code<br/>Player, Grid, Tree\"]\n"
    "        Assets[\"Asset Pixel & Audio<br/>sprites, rule-tiles, sfx\"]\n"
    "    end\n"
    "    \n"
    "    subgraph Storage_Layer [Storage]\n"
    "        Persistence[\"Persistence Data<br/>SaveData.json local\"]\n"
    "    end\n"
    "    \n"
    "    Unity <--> Gameplay\n"
    "    Unity <--> Assets\n"
    "    Gameplay <--> Persistence"
)
mermaid_run.font.name = 'Courier New'
mermaid_run.font.size = Pt(9.5)
mermaid_run.bold = True

spec_p = doc.add_paragraph()
spec_run = spec_p.add_run(
    "\nSPESIFIKASI TEKNIS:\n"
    "- Bahasa Pemrograman : C# (Object-Oriented Scripting)\n"
    "- Engine & Editor    : Unity Editor 2022.3+ / Unity 6\n"
    "- Pipeline Rendering : URP (Universal Render Pipeline)\n"
    "- Pembuatan Sprite   : Piskel Editor (Desain Ubin Pixel 16x16 / 32x32)\n"
    "- Data Penyimpanan   : Local JSON File (SaveData.json)"
)
spec_run.font.name = 'Arial'
spec_run.font.size = Pt(10)

doc.add_paragraph("\n" + "_"*50 + "\n")

# 4. Prototype Detail
doc.add_heading("4. Laporan Progres Prototipe Game (80% Selesai)", level=1)
p4 = doc.add_paragraph()
p4.add_run("Prototipe fungsional game telah dikembangkan langsung di dalam Unity dengan persentase kesiapan mencapai 80%. Berikut komponen core beserta cuplikan kode (code snippet) penting yang mendasari sistem tersebut:")

def add_code_block(doc_obj, code_text):
    p = doc_obj.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Add border and shading to code block
    pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                     r'<w:left w:val="single" w:sz="24" w:space="8" w:color="22703F"/>'
                     r'</w:pBdr>')
    shd = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F4F8F5"/>')
    p._p.get_or_add_pPr().append(pBdr)
    p._p.get_or_add_pPr().append(shd)
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(30, 70, 45)

proto_items = [
    ("Karakter & Fisika Gerak (PlayerController.cs)", 
     "Mendukung pergerakan 8 arah (WASD), melompat, merangkak, dan dash horizontal cepat. Dash menonaktifkan gravitasi sementara waktu untuk kontrol yang tajam.",
     "// Mekanisme Dash Celeste-Style pada PlayerController.cs\n"
     "private IEnumerator Dash() {\n"
     "    canDash = false; isDashing = true;\n"
     "    float originalGravity = rb.gravityScale;\n"
     "    rb.gravityScale = 0f; // Menghilangkan gravitasi sesaat\n"
     "    float dashDir = Input.GetAxisRaw(\"Horizontal\");\n"
     "    if (dashDir == 0) dashDir = Mathf.Sign(transform.localScale.x);\n"
     "    rb.linearVelocity = new Vector2(dashDir * dashSpeed, 0f);\n"
     "    yield return new WaitForSeconds(dashDuration);\n"
     "    rb.gravityScale = originalGravity; isDashing = false;\n"
     "    yield return new WaitForSeconds(dashCooldown);\n"
     "    canDash = true;\n"
     "}"),
    
    ("Sistem Gridworld & Sel Data (GridWorldMatrix.cs)", 
     "Menyimpan kelembaban tanah, kadar oksigen lokal, dan referensi objek di setiap koordinat peta menggunakan Dictionary dinamis.",
     "// Struktur Data Sel Grid & Penyimpanan Kamus pada GridWorldMatrix.cs\n"
     "[System.Serializable]\n"
     "public class GridCell {\n"
     "    [Range(0f, 1f)] public float moisture = 0.5f;\n"
     "    public float localO2 = 15.0f;\n"
     "    [Range(0f, 1f)] public float soilQuality = 0.5f;\n"
     "    public WorldObject placedObject;\n"
     "}\n"
     "public class GridWorldMatrix : MonoBehaviour {\n"
     "    private Dictionary<Vector2Int, GridCell> grid = new Dictionary<Vector2Int, GridCell>();\n"
     "    public GridCell GetCell(Vector2Int coordinates) {\n"
     "        if (!grid.ContainsKey(coordinates)) grid[coordinates] = new GridCell();\n"
     "        return grid[coordinates];\n"
     "    }\n"
     "}"),
     
    ("Logika Pertumbuhan Tanaman & FSM (Tree.cs)", 
     "Mengimplementasikan state machine tanaman (Seed -> Sprout -> Mature -> Withered). Tanaman memproduksi oksigen sesuai fase tumbuh dan mengonsumsi air tanah.",
     "// Siklus Hidup FSM & Pengecekan Kebutuhan Air pada Tree.cs\n"
     "public void ProgressGrowthCycle() {\n"
     "    if (currentFSMState == GrowthState.Withered) return;\n"
     "    if (ticksSinceLastWatered >= thresholdWaterRequirement) {\n"
     "        TransitionToWitheredState(); return;\n"
     "    }\n"
     "    switch (currentFSMState) {\n"
     "        case GrowthState.Seed: currentFSMState = GrowthState.Sprout; break;\n"
     "        case GrowthState.Sprout: currentFSMState = GrowthState.MatureTree; break;\n"
     "    }\n"
     "    ticksSinceLastWatered++; UpdateVisuals();\n"
     "}"),
     
    ("Pengaruh Atmosfer & Debuff Player (Player.cs)", 
     "Menghubungkan posisi pemain ke grid O2. Jika pemain masuk ke zona dengan oksigen rendah, kecepatannya otomatis melambat dan stamina berkurang sesuai rumus matematika debuff.",
     "// Evaluasi Debuff Kecepatan Berdasarkan Oksigen Koordinat pada Player.cs\n"
     "private void EvaluateCalculatedDebuffs() {\n"
     "    if (playerController == null) return;\n"
     "    Vector2Int playerGridPos = new Vector2Int(Mathf.RoundToInt(transform.position.x), Mathf.RoundToInt(-transform.position.y));\n"
     "    float localO2 = 15.0f;\n"
     "    if (EnvironmentManager.Instance != null && EnvironmentManager.Instance.EnvironmentGrid.HasCell(playerGridPos)) {\n"
     "        localO2 = EnvironmentManager.Instance.EnvironmentGrid.GetCell(playerGridPos).localO2;\n"
     "    }\n"
     "    float targetO2 = EnvironmentManager.Instance != null ? EnvironmentManager.Instance.TargetSafeO2 : 21.0f;\n"
     "    float o2Factor = Mathf.Min(1.0f, localO2 / targetO2);\n"
     "    playerController.moveSpeed = baseMovementSpeed * o2Factor; // Penerapan Debuff Speed\n"
     "}"),
     
    ("Map Loader & Import Aset (MapLoader.cs)", 
     "Membaca file data peta text (map_1.txt) dan menata ubin lantai, dinding pembatas, air, pasir, dan pohon secara instan pada Tilemap Unity Editor.",
     "// Parser File Teks Peta ke Unity Tilemap pada MapLoader.cs\n"
     "string[] rows = mapFile.text.Split(new[] { \"\\r\\n\", \"\\r\", \"\\n\" }, StringSplitOptions.RemoveEmptyEntries);\n"
     "for (int r = 0; r < rows.Length; r++) {\n"
     "    string[] cols = rows[r].Split(new[] { ' ' }, StringSplitOptions.RemoveEmptyEntries);\n"
     "    for (int c = 0; c < cols.Length; c++) {\n"
     "        if (int.TryParse(cols[c], out int tileIndex)) {\n"
     "            TileBase tileToSet = FindTileByIndex(tileIndex, indexToTileName);\n"
     "            if (tileToSet != null) tilemap.SetTile(new Vector3Int(c, -r, 0), tileToSet);\n"
     "        }\n"
     "    }\n"
     "}\n"
     "tilemap.RefreshAllTiles();"),

    ("Aset Grafis Pixel Art Kustom", 
     "Dibuat mandiri untuk ubin tanah baru, rumput hijau, batang pohon kering, serta ikon benih, tunas, dan pohon withered yang fungsional.",
     "// File Aset Gambar Kustom yang Diintegrasikan ke Proyek:\n"
     "// - Assets/Assets/tiles/dirt.png\n"
     "// - Assets/Assets/tiles/grass_0.png\n"
     "// - Assets/Assets/tiles/tree.png\n"
     "// - Assets/Assets/tiles/seed.png\n"
     "// - Assets/Assets/tiles/sprout.png\n"
     "// - Assets/Assets/tiles/withered.png")
]

for title, desc, code in proto_items:
    item_p = doc.add_paragraph()
    item_p.add_run(f"- {title}: ").bold = True
    item_p.add_run(desc)
    add_code_block(doc, code)

try:
    doc.save("Laporan_Konsep_Game_Life_on_Land_v2.docx")
    print("DOCX with snippets generated successfully as Laporan_Konsep_Game_Life_on_Land_v2.docx.")
except PermissionError:
    doc.save("Laporan_Konsep_Game_Life_on_Land_v3.docx")
    print("Warning: Laporan_Konsep_Game_Life_on_Land_v2.docx is locked. Saved as Laporan_Konsep_Game_Life_on_Land_v3.docx instead.")
